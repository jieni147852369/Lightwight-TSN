#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""补充“接口设计”章节。"""
import docx

DOC = "轻量时间敏感网络详细设计.docx"
d = docx.Document(DOC)


def find_idx(text):
    for i, p in enumerate(d.paragraphs):
        if p.text.strip() == text:
            return i
    raise ValueError(f"未找到标题: {text}")


def fill_section(next_title, blocks):
    anchor = d.paragraphs[find_idx(next_title)]
    for level, text in blocks:
        p = anchor.insert_paragraph_before(text)
        p.style = d.styles["Normal"] if level == 0 else d.styles[f"Heading {level}"]


IFACE = [
    (2, "帧头协议接口"),
    (0, "定义于 frame.h 的 struct frame_header：magic（0x54534e54=\"TSNT\" 业务帧，"
        "0x57524d55 预热帧）、header_size、seq（序号）、send_time_ns（发送 TAI 时间戳）、"
        "payload_len、priority。发送端填充，接收端据此校验并计算时延，是收发两端的核心数据契约。"),

    (2, "配置文件接口（INI）"),
    (0, "发送端 tsn_multi_sender.conf 与接收端 tsn_multi_receiver.conf 采用 INI 键值格式，"
        "由 config.h 解析装载：发送侧含 ifname、base_start、warmup_ms 及每帧的 name/dst/payload_len/"
        "period_us/socket_priority/udp_src_port/frame_limit（最多 MAX_FRAMES=16 帧，MAX_PAYLOAD=9000）；"
        "接收侧含 ifname、port。gptp.conf 提供 gPTP 主时钟参数。"),

    (2, "进程命令行接口"),
    (0, "tsn_multi_sender / tsn_multi_receiver：通过 -f <conf> 指定配置文件，另可用环境变量控制运行行为"
        "（如 CPU 亲和性设置、TSN_FORCE_HWTSTAMP 强制硬件时间戳）。"),
    (0, "ts_masterd：-f gptp.conf -i <iface> -m 启动 gPTP 主时钟；ts_ctl：SET GRANDMASTER_SETTINGS_NP "
        "等子命令控制运行期参数；clk_bridge：完成 PHC→系统时钟同步。"),

    (2, "日志解析接口"),
    (0, "GUI 通过 ProcessLineBuffer 按行解析 native 进程 stdout。发送日志：[name] frame N: prio=X "
        "start_ns sched_ns wake_late_ns send_ns；接收日志：seq= len= prio= src= expected_send_ns= "
        "rx_ts_ns= rx_ts_src= latency=。GUI 据此提取时延、抖动等指标并可视化与导出。"),

    (2, "方案序列化接口（JSON）"),
    (0, "SenderConfig / ReceiverConfig 的 Io 类负责方案的 JSON 读写，实现 GUI 配置的持久化与复用；"
        "GUI 依据方案生成上文各初始化/整形命令并驱动 native 进程，形成从方案编辑到实机执行的闭环。"),
]

fill_section("异常处理与可观测性设计", IFACE)
d.save(DOC)
print("接口设计 已写入，段落数：", len(d.paragraphs))