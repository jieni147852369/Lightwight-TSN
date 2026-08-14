#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""为核心章节的现有测试用例分组并规整层级。"""
import docx

DOC = "轻量时间敏感网络测试报告.docx"
d = docx.Document(DOC)


def find_idx(text):
    for i, p in enumerate(d.paragraphs):
        if p.text.strip() == text:
            return i
    raise ValueError(f"未找到: {text}")


def insert_h2_before(next_text, title):
    anchor = d.paragraphs[find_idx(next_text)]
    p = anchor.insert_paragraph_before(title)
    p.style = d.styles["Heading 2"]


def bold_case(text):
    """把用例编号段落加粗，作为用例小标题。"""
    idx = find_idx(text)
    p = d.paragraphs[idx]
    if p.runs:
        for r in p.runs:
            r.bold = True
    elif p.text:
        t = p.text
        p.text = ""
        r = p.add_run(t)
        r.bold = True


# 1) 插入四个分组 Heading 2（在各组第一条用例前）
insert_h2_before("4.1.1_(1):", "6.1 TAS 门控调度测试")
insert_h2_before("4.2.1_(1):", "6.2 CBS 队列限速测试")
insert_h2_before("4.2.3_(1):", "6.3 端到端通信时延测试")
insert_h2_before("基础功能运行配置", "6.4 gPTP 时间同步测试")

# 2) 用例编号加粗
for c in ["4.1.1_(1):", "4.1.1_(2):", "4.2.1_(1):", "4.2.1_(2):",
          "4.2.1_(3):", "4.2.1_(4):", "4.2.2_(1):", "4.2.3_(1):"]:
    bold_case(c)

# 3) 时间同步子项加粗
for c in ["基础功能运行配置", "1.时间同步启动", "2.时间同步功能", "3.时间同步角色切换"]:
    try:
        bold_case(c)
    except ValueError:
        pass

d.save(DOC)
print("分组与层级已规整，段落数：", len(d.paragraphs))