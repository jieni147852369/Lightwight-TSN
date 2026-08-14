#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import docx

DOC = "轻量时间敏感网络测试报告.docx"
d = docx.Document(DOC)
print("=== 段落总数:", len(d.paragraphs), "===")
for i, p in enumerate(d.paragraphs):
    t = p.text.strip()
    style = p.style.name if p.style else "?"
    if style.startswith("Heading"):
        print(f"{i:3d} | {style:12s} | {t[:60]}")