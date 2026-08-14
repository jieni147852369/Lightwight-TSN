#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""在封面之后、现有测试用例之前插入报告前置章节。"""
import docx

DOC = "轻量时间敏感网络测试报告.docx"
d = docx.Document(DOC)


def find_idx(text):
    for i, p in enumerate(d.paragraphs):
        if p.text.strip() == text:
            return i
    raise ValueError(f"未找到锚点: {text}")


def fill_before(next_text, blocks):
    anchor = d.paragraphs[find_idx(next_text)]
    for level, text in blocks:
        p = anchor.insert_paragraph_before(text)
        p.style = d.styles["Normal"] if level == 0 else d.styles[f"Heading {level}"]


FRONT = [
    (1, "1 概述"),
    (2, "1.1 编写目的"),
    (0, "本报告记录轻量时间敏感网络（TSN）软件的功能与性能测试过程及结果，验证时间感知整形（TAS）、"
        "信用整形（CBS）、端到端时延与 gPTP 时间同步等核心能力是否满足设计要求，为项目验收与后续优化提供依据。"),
    (2, "1.2 项目背景"),
    (0, "轻量时间敏感网络软件基于 Linux TSN 能力（tc taprio/mqprio/cbs、SO_TIMESTAMPING、CLOCK_TAI、"
        "IEEE 802.1AS gPTP）实现确定性以太网传输，提供发送、接收与时间同步的一体化配置与实验平台。"),

    (1, "2 测试依据与术语"),
    (2, "2.1 参考文档"),
    (0, "《轻量时间敏感网络详细设计》、IEEE 802.1Qbv（TAS）、IEEE 802.1Qav（CBS）、IEEE 802.1AS（gPTP）、"
        "Linux tc-taprio/tc-cbs 手册。"),
    (2, "2.2 术语与缩略语"),
    (0, "TSN：时间敏感网络；TAS：时间感知整形（门控调度）；CBS：信用整形；GCL：门控列表；"
        "gPTP：通用精确时间协议；PHC：PTP 硬件时钟；TAI：国际原子时；PCP：优先级代码点。"),

    (1, "3 被测对象与测试范围"),
    (2, "3.1 被测对象"),
    (0, "被测软件包含 GUI 配置端与 Native 执行端：GUI 负责方案编辑与 tc/ptp 命令构建；"
        "Native 端 tsn_multi_sender 负责有序实时发送、tsn_multi_receiver 负责接收与时间戳统计；"
        "时钟工具 ts_masterd/ts_ctl/clk_bridge 负责 gPTP 主时钟与 PHC→系统时钟同步。"),
    (2, "3.2 测试范围"),
    (0, "覆盖四类能力：TAS 门控调度、CBS 队列限速、端到端通信时延、gPTP 时间同步（含角色切换）；"
        "不含长稳老化测试与第三方交换机互操作测试。"),

    (1, "4 测试环境"),
    (2, "4.1 硬件环境"),
    (0, "三台主机构成主时钟、备用时钟、从时钟节点，均配置支持硬件时间戳（PHC）的网卡，通过千兆以太网互联。"),
    (2, "4.2 软件环境"),
    (0, "操作系统为支持 TSN 的 Linux 发行版，内核启用 tc taprio/mqprio/cbs 与 SO_TIMESTAMPING；"
        "运行被测软件的 GUI 与 Native 可执行物，时间同步使用 gptp.conf（hardware 时间戳、twoStep、"
        "network_transport L2、utc_offset 37）。"),
    (2, "4.3 网络拓扑与地址"),
    (0, "发送端向两台接收主机（192.168.150.2 / 192.168.150.3）发送多优先级流量；"
        "各优先级经 VLAN egress-qos-map（0:0..7:7）映射到对应硬件队列后由 TAS/CBS 调度。"),

  (1, "5 测试方案与用例设计"),
    (2, "5.1 测试方法"),
    (0, "采用实机黑盒测试：由 GUI 生成初始化与整形命令并下发，启动收发进程，通过发送端与接收端日志"
        "（帧序号、优先级、发送/接收时间戳、时延）核对实际行为与预期是否一致。"),
    (2, "5.2 用例分组"),
    (0, "用例按能力域编号：4.1.x 为 TAS 门控调度用例，4.2.1~4.2.3 为 CBS 限速与门控/时延用例，"
        "另设 gPTP 时间同步用例（启动、同步跟随、角色切换）。判定准则为“符合预期”即实际结果与设计目标一致。"),

    (1, "6 测试执行与结果"),
    (0, "本章为核心章节，逐条记录各用例的配置、发送与接收侧观测结果及结论。"),
]

fill_before("4.1.1_(1):", FRONT)
d.save(DOC)
print("前置章节已写入，段落数：", len(d.paragraphs))